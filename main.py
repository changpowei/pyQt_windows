#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt
from PyQt5.QtCore import *
import datetime
from News_scrawler import *
from Summarize import *
from os.path import exists

from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
import re
import datetime
import time
import sys
import os
import requests
import json
from opencc import OpenCC
import urllib
import ssl


class MyWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.content_save_location = './本文/'
        self.abstract_save_location = './摘要/'
        self.file_save_name = str(datetime.date.today())
        self.all_txt = ''
        self.initUI()


    def initUI(self):
        self.setWindowTitle('馬董早')
        self.setGeometry(50, 50, 800, 800)
        self.setFixedSize(800, 800)

        layout = QGridLayout()
        self.setLayout(layout)

        self.text_label = QLabel('請貼上文字', self)
        self.text_label.setStyleSheet("color:blue")
        self.text_label.setFont(QFont('Arial', 12))
        # self.text_label.setFixedSize()

        # 題目貼上位置
        self.text_input = QTextEdit()

        self.content_label = QLabel('內文：', self)
        self.content_label.setFont(QFont('Arial', 12))

        self.abstract_label = QLabel('摘要：', self)
        self.abstract_label.setFont(QFont('Arial', 12))

        self.content_folder_label = QLabel(self.content_save_location, self)
        self.content_folder_label.setStyleSheet("color:gray")
        self.content_folder_label.setFont(QFont('Arial', 12))
        self.content_folder_label.setStyleSheet("background-color: Gainsboro")
        self.content_folder_label.setFixedWidth(200)

        self.abstract_folder_label = QLabel(self.abstract_save_location, self)
        self.abstract_folder_label.setStyleSheet("color:gray")
        self.abstract_folder_label.setFont(QFont('Arial', 12))
        self.abstract_folder_label.setStyleSheet("background-color: Gainsboro")
        self.abstract_folder_label.setFixedWidth(200)

        self.output_file = QLabel('檔名：', self)
        self.output_file.setFont(QFont('Arial', 12))

        self.file_name = QLineEdit(self)
        self.file_name.setText(self.file_save_name)
        self.file_name.setAlignment(Qt.AlignCenter)


        self.content_output_btn = QPushButton('本文儲存位置', self)
        self.content_output_btn.clicked.connect(self.ContentButtonClick)

        self.abstract_output_btn = QPushButton('摘要儲存位置', self)
        self.abstract_output_btn.clicked.connect(self.AbstractButtonClick)

        self.crawler_btn = QPushButton('擷取全文', self)
        self.crawler_btn.clicked.connect(self.scratchButtonClick)
        self.crawler_btn.setEnabled(True)

        self.abstract_btn = QPushButton('擷取摘要', self)
        self.abstract_btn.clicked.connect(self.abstractButtonClick)
        self.abstract_btn.setEnabled(True)

        self.reset_btn = QPushButton('返回預設值', self)
        self.reset_btn.clicked.connect(self.resetButtonClick)

        self.show_infoes = QLabel(self)
        self.show_infoes.setStyleSheet("background-color: Gainsboro")
        self.show_infoes.setAlignment(Qt.AlignTop)
        self.show_infoes.setScaledContents(True)

        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)

        self.AllLayout(layout)

    def ContentButtonClick(self):

        folderPath = QtWidgets.QFileDialog.getExistingDirectory()  # 選取特定資料夾
        if folderPath == '':
            pass
        else:
            self.content_save_location = folderPath + '/'
            folder_split = folderPath.split('/')
            new_folder = '.../' + folder_split[-2] + '/' + folder_split[-1] + '/'
            self.content_folder_label.setText(new_folder)
            print(self.content_save_location)

    def AbstractButtonClick(self):

        folderPath = QtWidgets.QFileDialog.getExistingDirectory()  # 選取特定資料夾
        if folderPath == '':
            pass
        else:
            self.abstract_save_location = folderPath + '/'
            folder_split = folderPath.split('/')
            new_folder = '.../' + folder_split[-2] + '/' + folder_split[-1] + '/'
            self.abstract_folder_label.setText(new_folder)
            print(self.abstract_save_location)

    def scratchButtonClick(self):
        if self.abstract_btn.isEnabled() == False:
            QMessageBox.warning(None, '警告', '文章摘要執行中...')

        elif self.crawler_btn.isEnabled():
            if self.text_input.toPlainText() != "":
                self.crawler_btn.setDisabled(True)
                self.file_save_name = self.file_name.text()
                self.all_txt = self.text_input.toPlainText()

                self.news_crawler = News_scrawler(all_txt = self.all_txt, data=self.file_save_name, docx_save_loc=self.content_save_location)
                self.news_crawler.sinOut.connect(self.show_info)
                self.news_crawler.start()
            else:
                self.show_info("請於左方欄位輸入欲進行文章抓取的內容！")

    def abstractButtonClick(self):
        if self.crawler_btn.isEnabled() == False:
            QMessageBox.warning(None, '警告', '本文擷取中...')
        elif self.abstract_btn.isEnabled():
            if exists(self.content_save_location + self.file_save_name + '_本文.docx'):
                self.abstract_btn.setDisabled(True)

                self.summarize = Summarizer(data = self.file_save_name, load_path = self.content_save_location, save_path = self.abstract_save_location)
                self.summarize.sinOut.connect(self.show_info)
                self.summarize.start()
            else:
                self.show_info("本文不存在，請先執行爬蟲！")

    def show_info(self, infos):
        previous_text = self.show_infoes.text() + "\n"
        self.show_infoes.setText(previous_text + infos)
        if infos == "\n文章爬蟲完成":
            self.crawler_btn.setEnabled(True)

        if infos == "\n文章摘要完成":
            self.abstract_btn.setEnabled(True)
    def resetButtonClick(self):
        self.content_save_location = './本文/'
        self.content_folder_label.setText(self.content_save_location)
        self.abstract_save_location = './摘要/'
        self.abstract_folder_label.setText(self.abstract_save_location)
        self.file_save_name = str(datetime.date.today())
        self.file_name.setText(self.file_save_name)
        self.text_input.setText("")
        self.show_infoes.setText("")
        self.all_txt = self.text_input.toPlainText()


    def openFolder(self):
        folderPath = QtWidgets.QFileDialog.getExistingDirectory()  # 選取特定資料夾
        print(folderPath)

    def AllLayout(self, layout):
        layout.addWidget(self.text_label, 1, 0, 1, 2)
        layout.addWidget(self.content_label, 0, 3, 1, 2)
        layout.addWidget(self.abstract_label, 1, 3, 1, 2)
        layout.addWidget(self.content_folder_label, 0, 5, 1, 4)
        layout.addWidget(self.abstract_folder_label, 1, 5, 1, 4)
        layout.addWidget(self.content_output_btn, 0, 10, 1, 2)
        layout.addWidget(self.abstract_output_btn, 1, 10, 1, 2)
        layout.addWidget(self.text_input, 2, 0, 10, 8)
        layout.addWidget(self.output_file, 2, 8, 1, 2)
        layout.addWidget(self.file_name, 2, 10, 1, 2)
        layout.addWidget(self.crawler_btn, 10, 8, 1, 4)
        layout.addWidget(self.abstract_btn, 11, 8, 1, 4)
        layout.addWidget(self.show_infoes, 3, 8, 7, 4)
        layout.addWidget(self.reset_btn, 0, 0, 1, 2)
        layout.addWidget(self.scroll_area, 3, 8, 7, 4)

        self.scroll_area.setWidget(self.show_infoes)

class News_scrawler(QThread):
    sinOut = pyqtSignal(str)  # 自定義訊號，執行run()函式時，從相關執行緒發射此訊號
    def __init__(self, all_txt, data, docx_save_loc):
        super(News_scrawler, self).__init__()

        self.all_txt = all_txt  #要進行爬蟲的內容
        self.data = data    #檔案名稱
        self.docx_save_loc = docx_save_loc  #本文儲存位置

        if not os.path.isdir(self.docx_save_loc):
            os.mkdir(self.docx_save_loc)

        self.garbage_content = ['圖:', '資料照', '@', '▲', '圖/', '圖為', '取自',
                           '新頭殼newtalk', 'https',
                           '(臉書)', '(美聯社)', '(示意圖)', '(推特)',
                           '(左)', '(右)',
                           '作者:'
                           ]
        self.garbage_end = ['延伸閱讀', '德國之聲版權聲明', '《工商時報》LINE好友', 'LINE官方帳號', '立即加入NOWnews',
                       '更多鉅亨報導', '更多太報報導', '更多今周刊文章', '點我看更多華視新聞', '點我加入經濟日報好友',
                       '僅供參考', '推薦閱讀', '相關連結請見', '相關影片請見', '【NOWnews 今日新聞】提醒您'
                       ]

    def read_txt(self, path):
        text = []
        with open(path, 'r', encoding="utf-8") as file:
            for line in file.readlines():
                text.append(line)
                # print(line)
            return text

    def strQ2B(self, ustring):
        rstring = ""
        for uchar in ustring:
            inside_code = ord(uchar)
            if inside_code == 12288:                            # 全形空格直接轉換
                inside_code = 32
            elif 65281 <= inside_code <= 65374:   				# 全形字元（除空格）根據關係轉化
                inside_code -= 65248
            rstring += chr(inside_code)
        return rstring

    def crawler_web(self, url):

        options = Options()
        options.add_argument("--disable-notifications")

        chrome = webdriver.Chrome('./chromedriver', options=options)

        chrome.set_page_load_timeout(10)

        try:
            video_flag = False

            chrome.get(url)

            time.sleep(0.2)

            if 'money.udn.com' in url: #經濟日報
                content = chrome.find_element(By.ID, "article_body").text
            elif 'udn.com' in url:  #聯合報
                content = chrome.find_element(By.TAG_NAME, "article").text
            elif 'cna.com.tw' in url: #CNA
                content = chrome.find_element(By.CSS_SELECTOR, ".paragraph").text
            else:   #line
                if chrome.find_elements(By.XPATH, "//div[@class='videoEndContainer']"):
                    video_flag = True
                    collapse_button = chrome.find_element(By.CSS_SELECTOR, '.collapsiblePanel-toggleButton')
                    collapse_button.click()
                    time.sleep(0.1)
                    content = chrome.find_element(By.TAG_NAME, "article").text
                    content = '此題為影片:以下為影片概述\n' + content
                else:
                    content = chrome.find_element(By.TAG_NAME, "article").text

            content_splited = content.split('\n')

            if video_flag == False:
                content_filtered = self.garbage_filtering(content_splited)
            else:
                content_filtered = content_splited

            return content_filtered
        except:
            return "此題爬蟲抓不到，請手動抓取!"

    def garbage_filtering(self, content_splited):
        content_filtered = []
        for paragraph in content_splited:

            end_flag = False

            # 如果整段為英文則直接跳過
            if re.sub(r'[^\w\s]','',paragraph.replace(" ", "")).encode('utf-8').isalnum() == True:
                continue

            for end_item in self.garbage_end:
                if end_item in paragraph:
                    if end_item != '延伸閱讀' and end_item != '推薦閱讀':
                        end_flag = True
                        break
                    else:
                        if len(paragraph) < 7:
                            end_flag = True
                            break

            if end_flag == True:
                break
            else:
                if '\u3000' in paragraph:
                    continue
                else:
                    garbage_flag = False
                    content_half = self.strQ2B(paragraph)
                    for garbage_item in self.garbage_content:
                        if garbage_item in content_half.replace(" ", ""):
                            garbage_flag = True
                            break
                    if garbage_flag == False:
                        content_filtered.append(content_half)
        return content_filtered


    def run(self):

        document = Document()

        # org_path = './解析txt/' + self.data + '.txt'
        save_path = self.docx_save_loc + self.data + '_本文.docx'

        # all_txt = self.read_txt(org_path)
        all_txt = self.all_txt.split("\n")

        questions = 0
        for line in all_txt:
            if line[:5] == 'https':
                questions += 1

        start = datetime.datetime.now()

        num = 0
        for line in all_txt:
            if(line != '' and line[0]!="共"):
                if(line[:5] == 'https'):
                    num += 1
                    content_filtered = self.crawler_web(line)
                    if isinstance(content_filtered, list):
                        for content in content_filtered:
                            document.add_paragraph()
                            document.add_paragraph(content)
                    else:
                        document.add_paragraph(content_filtered)
                        document.add_paragraph()
                        document.add_paragraph(line)

                    document.add_page_break()

                    print("\r", end="")
                    print("Crawl progress: {}%: ".format(int(num * 100 / questions)),
                          "▋" * (int(num * 100 / questions) // 2), end="")
                    sys.stdout.flush()
                    self.sinOut.emit("Progress: {}%".format(int(num * 100 / questions)))
                else:
                    document.add_heading(line, level=1)


        document.save(save_path)

        end = datetime.datetime.now()

        self.sinOut.emit("\n{} 題\n 執行時間 : {}".format(questions,  end - start))
        self.sinOut.emit("\n文章爬蟲完成")

        # return questions, (end - start)

        # print("\n")
        # print("{} 題, 執行時間 : {}".format(questions,  end - start))

class Summarizer(QThread):
    sinOut = pyqtSignal(str)  # 自定義訊號，執行run()函式時，從相關執行緒發射此訊號
    def __init__(self, data, load_path, save_path):
        super(Summarizer, self).__init__()

        if not os.path.isdir(save_path):
            os.mkdir(save_path)

        self.data = data
        self.load_path = load_path + self.data + '_本文.docx'
        self.save_path = save_path + self.data + '_摘要.docx'

    def TW2S(self, article):
        # 繁體轉簡體
        cc = OpenCC('tw2s')
        return cc.convert(article)

    def S2TW(self, article):
        # 簡體轉繁體
        cc = OpenCC('s2tw')
        return cc.convert(article)


    def Summarizing_free(self, one_article):
        try:
            token = 'grv8ObC13pwUfffZ'
            taskID = 'zhai0904'

            payload = {'token': token, 'taskid': taskID, 'con':one_article}
            Summarized = requests.post("http://inter.xiaohuaerai.com/api/task", data=payload)
            # print(json.loads(Summarized.text)["data"]["contents"])

            return json.loads(Summarized.text)["data"]["contents"]

        except:

            return "文章無法抽取摘要，請自行剪貼!"

    def Summarizing_paid(self, one_article):
        try:
            host = 'https://zhaiyao.xiaohuaerai.com'
            path = '/zhaiyao'
            method = 'POST'
            appcode = '449f24b2833844c6a5b0b51fed1b6f3e'
            querys = ''
            bodys = {}
            url = host + path

            bodys['src'] = one_article
            post_data = urllib.parse.urlencode(bodys).encode("utf-8")
            request = urllib.request.Request(url, post_data)
            request.add_header('Authorization', 'APPCODE ' + appcode)
            # 根据API的要求，定义相对应的Content - Type
            request.add_header('Content-Type', 'application/x-www-form-urlencoded; charset=UTF-8')
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            response = urllib.request.urlopen(request, context=ctx)
            content = response.read()
            # if (content):
            #     print(json.loads(content)["res"])
            return json.loads(content)["res"]

        except Exception as e :

            print(e)
            return "文章無法抽取摘要，請自行剪貼!"

    def GetSummarize(self, load_path, save_path):

        load_content = Document(load_path)
        summarized_content = Document()

        one_article = ""
        question_nums = 0
        for i, paragraph in enumerate(load_content.paragraphs):

            if paragraph.style.name == 'Heading 1':
                question_nums += 1
                summarized_content.add_heading(paragraph.text, level=1)
                self.sinOut.emit("第{}題: {}".format(question_nums, paragraph.text))
                print("第{}題: {}".format(question_nums, paragraph.text))

                one_article = ""

            else:
                if paragraph.text != "" and paragraph.text != "\n":
                        one_article += paragraph.text


                if i == len(load_content.paragraphs)-1:
                    if one_article != []:
                        self.Transfering(summarized_content, one_article)

                elif  load_content.paragraphs[i+1].style.name == 'Heading 1':
                    self.Transfering(summarized_content, one_article)

        summarized_content.save(save_path)

    def Transfering(self, summarized_content, one_article):

        summarized_content.add_paragraph(self.Summarizing_free(one_article))
        summarized_content.add_paragraph()
        summarized_content.add_page_break()


    def run(self):

        # data = 'test'

        # load_path = './本文/' + self.data + '_本文.docx'
        # save_path = './摘要/' + self.data + '_摘要.docx'

        self.GetSummarize(self.load_path, self.save_path)

        self.sinOut.emit("\n文章摘要完成")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    w = MyWidget()
    w.show()
    sys.exit(app.exec_())