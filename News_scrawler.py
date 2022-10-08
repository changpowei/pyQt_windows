from gc import garbage
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
import re
import datetime
import time
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
import sys
import os


class News_scrawler(QThread):
    sinOut = pyqtSignal(str)  # 自定義訊號，執行run()函式時，從相關執行緒發射此訊號
    def __init__(self, all_txt, data, docx_save_loc):
        super(News_scrawler, self).__init__()

        self.all_txt = all_txt  #要進行爬蟲的內容
        self.data = data    #檔案名稱
        self.docx_save_loc = docx_save_loc  #本文儲存位置

        if not os.path.isdir(self.docx_save_loc):
            os.mkdir(self.docx_save_loc)

        self.garbage_content = ['圖:', '資料照', '@', '▲', '圖/', '圖為', '取自',
                           '新頭殼newtalk', 'https',
                           '(臉書)', '(美聯社)', '(示意圖)', '(推特)',
                           '(左)', '(右)',
                           '作者:'
                           ]
        self.garbage_end = ['延伸閱讀', '德國之聲版權聲明', '《工商時報》LINE好友', 'LINE官方帳號', '立即加入NOWnews',
                       '更多鉅亨報導', '更多太報報導', '更多今周刊文章', '點我看更多華視新聞', '點我加入經濟日報好友',
                       '僅供參考', '推薦閱讀', '相關連結請見', '相關影片請見', '【NOWnews 今日新聞】提醒您'
                       ]

    def read_txt(self, path):
        text = []
        with open(path, 'r', encoding="utf-8") as file:
            for line in file.readlines():
                text.append(line)
                # print(line)
            return text

    def strQ2B(self, ustring):
        rstring = ""
        for uchar in ustring:
            inside_code = ord(uchar)
            if inside_code == 12288:                            # 全形空格直接轉換
                inside_code = 32
            elif 65281 <= inside_code <= 65374:   				# 全形字元（除空格）根據關係轉化
                inside_code -= 65248
            rstring += chr(inside_code)
        return rstring

    def crawler_web(self, url):

        options = Options()
        options.add_argument("--disable-notifications")

        chrome = webdriver.Chrome('./chromedriver', options=options)

        chrome.set_page_load_timeout(10)

        try:
            video_flag = False

            chrome.get(url)

            time.sleep(0.2)

            if 'money.udn.com' in url: #經濟日報
                content = chrome.find_element(By.ID, "article_body").text
            elif 'udn.com' in url:  #聯合報
                content = chrome.find_element(By.TAG_NAME, "article").text
            elif 'cna.com.tw' in url: #CNA
                content = chrome.find_element(By.CSS_SELECTOR, ".paragraph").text
            else:   #line
                if chrome.find_elements(By.XPATH, "//div[@class='videoEndContainer']"):
                    video_flag = True
                    collapse_button = chrome.find_element(By.CSS_SELECTOR, '.collapsiblePanel-toggleButton')
                    collapse_button.click()
                    time.sleep(0.1)
                    content = chrome.find_element(By.TAG_NAME, "article").text
                    content = '此題為影片:以下為影片概述\n' + content
                else:
                    content = chrome.find_element(By.TAG_NAME, "article").text

            content_splited = content.split('\n')

            if video_flag == False:
                content_filtered = self.garbage_filtering(content_splited)
            else:
                content_filtered = content_splited

            return content_filtered
        except:
            return "此題爬蟲抓不到，請手動抓取!"

    def garbage_filtering(self, content_splited):
        content_filtered = []
        for paragraph in content_splited:

            end_flag = False

            # 如果整段為英文則直接跳過
            if re.sub(r'[^\w\s]','',paragraph.replace(" ", "")).encode('utf-8').isalnum() == True:
                continue

            for end_item in self.garbage_end:
                if end_item in paragraph:
                    if end_item != '延伸閱讀' and end_item != '推薦閱讀':
                        end_flag = True
                        break
                    else:
                        if len(paragraph) < 7:
                            end_flag = True
                            break

            if end_flag == True:
                break
            else:
                if '\u3000' in paragraph:
                    continue
                else:
                    garbage_flag = False
                    content_half = self.strQ2B(paragraph)
                    for garbage_item in self.garbage_content:
                        if garbage_item in content_half.replace(" ", ""):
                            garbage_flag = True
                            break
                    if garbage_flag == False:
                        content_filtered.append(content_half)
        return content_filtered


    def run(self):

        document = Document()

        # org_path = './解析txt/' + self.data + '.txt'
        save_path = self.docx_save_loc + self.data + '_本文.docx'

        # all_txt = self.read_txt(org_path)
        all_txt = self.all_txt.split("\n")

        questions = 0
        for line in all_txt:
            if line[:5] == 'https':
                questions += 1

        start = datetime.datetime.now()

        num = 0
        for line in all_txt:
            if(line != '' and line[0]!="共"):
                if(line[:5] == 'https'):
                    num += 1
                    content_filtered = self.crawler_web(line)
                    if isinstance(content_filtered, list):
                        for content in content_filtered:
                            document.add_paragraph()
                            document.add_paragraph(content)
                    else:
                        document.add_paragraph(content_filtered)
                        document.add_paragraph()
                        document.add_paragraph(line)

                    document.add_page_break()

                    print("\r", end="")
                    print("Crawl progress: {}%: ".format(int(num * 100 / questions)),
                          "▋" * (int(num * 100 / questions) // 2), end="")
                    sys.stdout.flush()
                    self.sinOut.emit("Progress: {}%".format(int(num * 100 / questions)))
                else:
                    document.add_heading(line, level=1)


        document.save(save_path)

        end = datetime.datetime.now()

        self.sinOut.emit("\n{} 題\n 執行時間 : {}".format(questions,  end - start))
        self.sinOut.emit("\n文章爬蟲完成")

        # return questions, (end - start)

        # print("\n")
        # print("{} 題, 執行時間 : {}".format(questions,  end - start))

