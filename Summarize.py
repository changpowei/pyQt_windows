#-*- coding: utf-8 -*-

from lib2to3.pgen2.token import tok_name
from docx import Document
import requests
import json
import base64
from opencc import OpenCC
import urllib
import ssl
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
import os

class Summarizer(QThread):
    sinOut = pyqtSignal(str)  # 自定義訊號，執行run()函式時，從相關執行緒發射此訊號
    def __init__(self, data, load_path, save_path):
        super(Summarizer, self).__init__()

        if not os.path.isdir(save_path):
            os.mkdir(save_path)

        self.data = data
        self.load_path = load_path + self.data + '_本文.docx'
        self.save_path = save_path + self.data + '_摘要.docx'

    def TW2S(self, article):
        # 繁體轉簡體
        cc = OpenCC('tw2s')
        return cc.convert(article)

    def S2TW(self, article):
        # 簡體轉繁體
        cc = OpenCC('s2tw')
        return cc.convert(article)


    def Summarizing_free(self, one_article):
        try:
            token = 'grv8ObC13pwUfffZ'
            taskID = 'zhai0904'

            payload = {'token': token, 'taskid': taskID, 'con':one_article}
            Summarized = requests.post("http://inter.xiaohuaerai.com/api/task", data=payload)
            # print(json.loads(Summarized.text)["data"]["contents"])

            return json.loads(Summarized.text)["data"]["contents"]

        except:

            return "文章無法抽取摘要，請自行剪貼!"

    def Summarizing_paid(self, one_article):
        try:
            host = 'https://zhaiyao.xiaohuaerai.com'
            path = '/zhaiyao'
            method = 'POST'
            appcode = '449f24b2833844c6a5b0b51fed1b6f3e'
            querys = ''
            bodys = {}
            url = host + path

            bodys['src'] = one_article
            post_data = urllib.parse.urlencode(bodys).encode("utf-8")
            request = urllib.request.Request(url, post_data)
            request.add_header('Authorization', 'APPCODE ' + appcode)
            # 根据API的要求，定义相对应的Content - Type
            request.add_header('Content-Type', 'application/x-www-form-urlencoded; charset=UTF-8')
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            response = urllib.request.urlopen(request, context=ctx)
            content = response.read()
            # if (content):
            #     print(json.loads(content)["res"])
            return json.loads(content)["res"]

        except Exception as e :

            print(e)
            return "文章無法抽取摘要，請自行剪貼!"

    def GetSummarize(self, load_path, save_path):

        load_content = Document(load_path)
        summarized_content = Document()

        one_article = ""
        question_nums = 0
        for i, paragraph in enumerate(load_content.paragraphs):

            if paragraph.style.name == 'Heading 1':
                question_nums += 1
                summarized_content.add_heading(paragraph.text, level=1)
                self.sinOut.emit("第{}題: {}".format(question_nums, paragraph.text))
                print("第{}題: {}".format(question_nums, paragraph.text))

                one_article = ""

            else:
                if paragraph.text != "" and paragraph.text != "\n":
                        one_article += paragraph.text


                if i == len(load_content.paragraphs)-1:
                    if one_article != []:
                        self.Transfering(summarized_content, one_article)

                elif  load_content.paragraphs[i+1].style.name == 'Heading 1':
                    self.Transfering(summarized_content, one_article)

        summarized_content.save(save_path)

    def Transfering(self, summarized_content, one_article):

        summarized_content.add_paragraph(self.Summarizing_free(one_article))
        summarized_content.add_paragraph()
        summarized_content.add_page_break()


    def run(self):

        # data = 'test'

        # load_path = './本文/' + self.data + '_本文.docx'
        # save_path = './摘要/' + self.data + '_摘要.docx'

        self.GetSummarize(self.load_path, self.save_path)

        self.sinOut.emit("\n文章摘要完成")
